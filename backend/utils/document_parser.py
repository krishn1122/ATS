import PyPDF2
import docx
import re
import logging
from typing import Optional, Dict, List
from io import BytesIO
import tempfile
import os

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentParser:
    """Enhanced document parser supporting PDF and DOCX formats with error handling"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> Optional[str]:
        """Extract text from PDF file with enhanced error handling"""
        try:
            pdf_stream = BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_stream)
            
            text = ""
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                    continue
            
            return text.strip() if text.strip() else None
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            return None
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> Optional[str]:
        """Extract text from DOCX file with enhanced error handling"""
        try:
            # Create a temporary file to store the content
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = docx.Document(temp_file_path)
                text_parts = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text.strip())
                
                text = "\n".join(text_parts)
                return text.strip() if text.strip() else None
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return None
    
    @staticmethod
    def parse_document(file_content: bytes, file_type: str) -> Optional[str]:
        """Parse document based on file type"""
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return DocumentParser.extract_text_from_pdf(file_content)
        elif file_type == 'docx':
            return DocumentParser.extract_text_from_docx(file_content)
        else:
            logger.error(f"Unsupported file type: {file_type}")
            return None

class TextAnalyzer:
    """Text analysis utilities for grammar, repetition, and format checking"""
    
    @staticmethod
    def analyze_grammar(text: str) -> List[Dict]:
        """Basic grammar analysis - can be enhanced with external libraries"""
        issues = []
        lines = text.split('\n')
        
        # Basic grammar patterns to check
        patterns = {
            'double_space': r'  +',
            'missing_capitalization': r'\. [a-z]',
            'missing_period': r'[a-zA-Z]\n[A-Z]',
        }
        
        for line_num, line in enumerate(lines, 1):
            for issue_type, pattern in patterns.items():
                matches = re.finditer(pattern, line)
                for match in matches:
                    issues.append({
                        'text': match.group(),
                        'line_number': line_num,
                        'suggestion': TextAnalyzer._get_grammar_suggestion(issue_type),
                        'severity': 'medium'
                    })
        
        return issues
    
    @staticmethod
    def _get_grammar_suggestion(issue_type: str) -> str:
        """Get grammar suggestions based on issue type"""
        suggestions = {
            'double_space': 'Use single space between words',
            'missing_capitalization': 'Capitalize first letter after period',
            'missing_period': 'Add period at end of sentence',
        }
        return suggestions.get(issue_type, 'Check grammar')
    
    @staticmethod
    def analyze_word_repetition(text: str) -> List[Dict]:
        """Analyze word repetition in the text"""
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        word_count = {}
        word_positions = {}
        
        for i, word in enumerate(words):
            if word not in word_count:
                word_count[word] = 0
                word_positions[word] = []
            word_count[word] += 1
            word_positions[word].append(i)
        
        # Find words repeated more than reasonable threshold
        repetitions = []
        for word, count in word_count.items():
            if count > 3 and len(word) > 4:  # Configurable thresholds
                repetitions.append({
                    'word': word,
                    'count': count,
                    'positions': word_positions[word]
                })
        
        return sorted(repetitions, key=lambda x: x['count'], reverse=True)
    
    @staticmethod
    def analyze_format_issues(text: str) -> List[Dict]:
        """Analyze format and readability issues"""
        issues = []
        lines = text.split('\n')
        
        # Check for inconsistent formatting
        if len([line for line in lines if line.strip()]) < len(lines) * 0.7:
            issues.append({
                'issue_type': 'spacing',
                'description': 'Too many empty lines affecting readability',
                'suggestion': 'Reduce excessive blank lines'
            })
        
        # Check for very long lines
        long_lines = [line for line in lines if len(line) > 100]
        if len(long_lines) > len(lines) * 0.3:
            issues.append({
                'issue_type': 'line_length',
                'description': 'Many lines are too long',
                'suggestion': 'Break long lines for better readability'
            })
        
        # Check for bullet point consistency
        bullet_patterns = [r'^\s*[-*•]\s', r'^\s*\d+\.\s', r'^\s*[a-zA-Z]\.\s']
        bullet_usage = sum(len(re.findall(pattern, text)) for pattern in bullet_patterns)
        if bullet_usage < 3 and len(text) > 500:
            issues.append({
                'issue_type': 'structure',
                'description': 'Consider using bullet points for better structure',
                'suggestion': 'Add bullet points for achievements and skills'
            })
        
        return issues
    
    @staticmethod
    def calculate_readability_score(text: str) -> float:
        """Calculate a basic readability score (0-100) optimized for resumes"""
        if not text.strip():
            return 0.0
        
        # Basic metrics
        sentences = len(re.findall(r'[.!?]+', text))
        words = len(re.findall(r'\b\w+\b', text))
        characters = len(re.sub(r'\s', '', text))
        
        if sentences == 0 or words == 0:
            return 65.0  # Default reasonable score for bullet-heavy resumes
        
        # Simple readability calculation (modified for resumes)
        avg_sentence_length = words / sentences
        avg_word_length = characters / words
        
        # Resume-optimized scoring (accounts for bullet points and technical terms)
        if avg_sentence_length < 5:  # Likely bullet points
            base_score = 75  # Good for ATS
        elif avg_sentence_length < 15:
            base_score = 70
        elif avg_sentence_length < 25:
            base_score = 60
        else:
            base_score = 45
        
        # Adjust for word complexity (technical resumes often have longer words)
        if avg_word_length < 4.5:
            word_bonus = 10
        elif avg_word_length < 6:
            word_bonus = 0
        else:
            word_bonus = -5  # Technical terms are acceptable in resumes
        
        final_score = max(35, min(90, base_score + word_bonus))
        
        return round(final_score, 2)